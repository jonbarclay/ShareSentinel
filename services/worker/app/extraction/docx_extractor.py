"""Word document (.docx) text extraction using python-docx."""

from pathlib import Path

from docx import Document

from .base import BaseExtractor, ExtractionResult


class DocxExtractor(BaseExtractor):
    """Extract text from Word documents including properties, body, tables, headers/footers."""

    def extract(self, file_path: Path, file_size: int) -> ExtractionResult:
        try:
            doc = Document(str(file_path))

            # Document properties
            props = doc.core_properties
            metadata = {
                "title": props.title or "",
                "author": props.author or "",
                "subject": props.subject or "",
                "keywords": props.keywords or "",
                "comments": props.comments or "",
                "last_modified_by": props.last_modified_by or "",
                "category": props.category or "",
            }

            parts: list[str] = []

            # Document properties as context
            prop_text: list[str] = []
            if metadata["title"]:
                prop_text.append(f"Document Title: {metadata['title']}")
            if metadata["author"]:
                prop_text.append(f"Author: {metadata['author']}")
            if metadata["subject"]:
                prop_text.append(f"Subject: {metadata['subject']}")
            if metadata["keywords"]:
                prop_text.append(f"Keywords: {metadata['keywords']}")
            if prop_text:
                parts.append("--- Document Properties ---\n" + "\n".join(prop_text))

            # Body text
            body_paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            if body_paragraphs:
                parts.append("--- Document Body ---\n" + "\n".join(body_paragraphs))

            # Table content
            for i, table in enumerate(doc.tables):
                rows_text: list[str] = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows_text.append(" | ".join(cells))
                if rows_text:
                    parts.append(f"--- Table {i + 1} ---\n" + "\n".join(rows_text))

            # Headers and footers
            header_footer_text: list[str] = []
            for section in doc.sections:
                if section.header and section.header.paragraphs:
                    for p in section.header.paragraphs:
                        if p.text.strip():
                            header_footer_text.append(f"Header: {p.text.strip()}")
                if section.footer and section.footer.paragraphs:
                    for p in section.footer.paragraphs:
                        if p.text.strip():
                            header_footer_text.append(f"Footer: {p.text.strip()}")
            if header_footer_text:
                parts.append(
                    "--- Headers/Footers ---\n" + "\n".join(header_footer_text)
                )

            full_text = "\n\n".join(parts)

            if len(full_text.strip()) < self.MIN_MEANINGFUL_TEXT:
                return ExtractionResult(
                    success=False,
                    extraction_method="docx_text",
                    metadata=metadata,
                    error="Document appears to be empty or contains only non-text elements",
                )

            text, was_sampled, description = self._truncate_if_needed(
                full_text,
                "First ~100KB of text from Word document",
            )

            return ExtractionResult(
                success=True,
                extraction_method="docx_text",
                text_content=text,
                metadata=metadata,
                was_sampled=was_sampled,
                sampling_description=description,
                content_length=len(text),
            )
        except Exception as e:
            return ExtractionResult(
                success=False,
                extraction_method="docx_text",
                error=str(e),
            )
